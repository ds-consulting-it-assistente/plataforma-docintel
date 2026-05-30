import os
import json
import streamlit as st
from pypdf import PdfReader
from docx import Document
import pandas as pd
from groq import Groq

# Inicializa o cliente Groq buscando a chave de forma segura das Configurações do Streamlit
def get_groq_client():
    api_key = st.secrets.get("GROQ_API_KEY") or os.environ.get("GROQ_API_KEY")
    if not api_key:
        st.error("Chave API do Groq não encontrada! Configure-a nas Secrets do Streamlit.")
        st.stop()
    return Groq(api_key=api_key)

def extract_text_from_pdf(pdf_file):
    try:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        return f"Erro ao ler PDF: {str(e)}"

def extract_text_from_docx(docx_file):
    try:
        doc = Document(docx_file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Erro ao ler Word: {str(e)}"

def read_universal_file(uploaded_file):
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    if name.endswith('.pdf'):
        return extract_text_from_pdf(uploaded_file)
    elif name.endswith('.docx'):
        return extract_text_from_docx(uploaded_file)
    elif name.endswith('.txt'):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    return ""

def ask_groq(system_prompt, user_content):
    client = get_groq_client()
    try:
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            temperature=0.3,
        )
        return completion.choices[0].message.content
    except Exception as e:
        return f"Erro na API Groq: {str(e)}"

def create_word_report(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("Relatório gerado automaticamente pela plataforma DocIntel AI.\n")
    
    # Divide por linhas para injetar no Word de forma limpa
    for line in content.split('\n'):
        if line.strip().startswith('#'):
            doc.add_heading(line.replace('#', '').strip(), level=2)
        else:
            doc.add_paragraph(line)
            
    from io import BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
